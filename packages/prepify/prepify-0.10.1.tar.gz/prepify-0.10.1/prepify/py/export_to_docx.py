from pathlib import Path

from docx import Document
import PySimpleGUI as sg

import prepify.py.edit_settings as es
import prepify.py.read_doc_form as rdf

####################
# Available Styles #
####################

def add_h1(document: Document, value: str):
    if value in [[], '', None]:
        return
    return document.add_heading(value, level=1)

def add_h2(document: Document, value: str):
    if value in [[], '', None]:
        return
    return document.add_heading(value, level=2)

def add_0ind(document: Document, text: str):
    return document.add_paragraph(text, style='Normal')

def add_1ind(document: Document, text: str):
    return document.add_paragraph(text, style='1indent')

def add_2ind(document: Document, text: str):
    return document.add_paragraph(text, style='2indent')

##############
# Formatting #
##############

def format_foliation(pd: dict):
    sections = []
    covers = []
    numbers = []
    for x in pd['foliation']:
        print(x)
        text = x.split('|')[1]
        if 'Cover' in text:
            # covers.append(f"{text.replace('Cover: ', '').lower().strip()} (inside & outside)")
            continue
        numbers.append(x.split(':')[1].split()[0])
        sections.append(text.replace('(', '(Written numbers '))
    return sections, covers, numbers

def format_quire_structure(pd: dict):
    quires = []
    for x in pd['quire_structure']:
        quires.append(
            x.split(':')[1].strip()
        )
    return quires

def format_covers(document, covers: list, pd: dict):
    if pd['cover_edge']:
        covers.append('edge')
    if pd['cover_spine']:
        covers.append('spine')
    if pd['cover_top']:
        covers.append('top')
    if pd['cover_bottom']:
        covers.append('bottom')
    return covers

##################
# Write Document #
##################

def title(document: Document, pd: dict):
    if pd['ga_number'] != '':
        document.add_heading(f"Gregory-Aland {pd['ga_number']}", 0)
    else:
        document.add_heading(pd['csntm_id'], 0)
    subtitle = f'({pd["institution"]}: {pd["shelf_number"]})'
    document.add_paragraph(subtitle, style='Subtitle')

def description(document: Document, pd: dict):
    # get data
    ga = pd['ga_number'] if pd['ga_number'] != '' else '—'
    ga = f'GA Number: {ga}'
    contents = f'Contents: {rdf.format_content(pd)}'
    date = f'Date: {rdf.format_date(pd)}'
    material = f'Material: {rdf.format_material(pd)}'
    leaves = f"Leaves: {pd.get('number_of_leaves', '—')}"
    columns = f"Columns: {pd.get('number_of_columns', '—')}"
    lines = f"Lines per page: {rdf.format_lines(pd)}"
    shelf_number = f"Shelf Number: {pd.get('shelf_number', '—')}"
    width, height, depth = rdf.parse_dimensions(pd)
    dimensions = f'Dimensions: {width} W x {height} H x {depth}'
    # print data
    add_h1(document, 'CSNTM description:')
    add_0ind(document, ga)
    add_0ind(document, contents)
    add_0ind(document, date)
    add_0ind(document, material)
    add_0ind(document, leaves)
    add_0ind(document, columns)
    add_0ind(document, lines)
    add_0ind(document, shelf_number)
    add_0ind(document, dimensions)

def codicological_observations(pd: dict, document: Document):
    # get data
    folios, covers, numbers = format_foliation(pd)
    total_folios = 0
    for num in numbers:
        try:
            total_folios += int(num)
        except:
            print(f'Not a number: {num}')
    covers = format_covers(document, covers, pd)
    cover_images = 0
    for c in covers:
        if c.startswith('front') or c.startswith('back'):
            cover_images += 2
        else:
            cover_images += 1
    covers = ', '.join(covers)
    quires = format_quire_structure(pd)    
    # print data
    add_0ind(document, f'Total Images: {cover_images+(total_folios*2)}')
    add_h1(document, 'Codicological Observations')
    if folios != []:
        foliation = add_h2(document, 'Foliation: ')
        foliation.add_run(f'({total_folios} total folios/{total_folios*2} images)').bold = False
        for section in folios:
            add_2ind(document, section.strip())
    if quires != []:
        add_h2(document, 'Quires:')
        for quire in quires:
            add_2ind(document, quire)
    if covers != []:
        cover_form = add_h2(document, 'Covers: ')
        cover_form.add_run(f'({cover_images} total images)').bold = False
        add_2ind(document, covers)

def notable_features(document: Document, pd: dict):
    add_h1(document, 'Notable features:')
    if pd['table_of_contents'] != []:
        add_h2(document, 'Table of Contents:')
        for item in pd['table_of_contents']:
            add_2ind(document, item)

########
# MAIN #
########

def to_docx(pd: dict):
    settings = es.get_settings()
    main_dir = Path(__file__).parent.parent.as_posix()
    document = Document(docx=f'{main_dir}/prepdoc_template.docx')
    
    title(document, pd)
    description(document, pd)
    codicological_observations(pd, document)
    notable_features(document, pd)
    new_filename = sg.popup_get_file(
        '', save_as=True, 
        file_types=(('Word Document', '*.docx'),), 
        initial_folder=settings.get('project_folder'),
        no_window=True
        )
    if not new_filename:
        return
    try:
        document.save(new_filename)
    except PermissionError:
        sg.popup_ok('Cannot change an opened DOCX file.\nClose it first or save this one with a different file name.', title='DOCX File Already Open')
    try:
        es.update_settings('project_folder', Path(new_filename).parent.as_posix())
    except:
        print('failed to update settings')
